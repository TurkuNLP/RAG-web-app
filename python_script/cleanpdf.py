import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt
import os
import re
from PIL import Image
import pytesseract
import io
from docxcompose.composer import Composer

def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    text = ""
    for page_num in range(len(doc)):

        page = doc.load_page(page_num)
        page_text = page.get_text("text")

        # if it is an image or pdf belonging before 2007
        if pdf_path.split('/')[-1].find('2003') == 0 or \
            pdf_path.split('/')[-1].find('2002') == 0 or \
            pdf_path.split('/')[-1].find('2001') == 0:
            
            image = page.get_pixmap()
            pix = page.get_pixmap(matrix=image)
        
            # Convert the pixmap object to a PIL image object
            img = Image.open(io.BytesIO(pix.tobytes()))
        
            # Use pytesseract to do OCR on the image
            page_text = pytesseract.image_to_string(img, lang='rus+eng+deu')
            
        print(pdf_path.split('/')[-1], page_num)
        lines = page_text.split('\n')
        
        paragraph = ""
        for line in lines:
            stripped_line = line.strip()
            if stripped_line:
                if paragraph:
                    if paragraph.endswith('-'):
                        paragraph = paragraph[:-1] + stripped_line  # Reconstituer le mot coupé
                    else:
                        paragraph += " " + stripped_line
                else:
                    paragraph = stripped_line
            else:
                if paragraph:
                    text += paragraph + "\n\n"
                    paragraph = ""
        
        if paragraph:
            text += paragraph + "\n\n"
        
        # Ajouter le numéro de page à la fin de chaque page
        text += f"Page {page_num + 1}\n\n"

    return text.strip()

def clean_text(text):
    # Suppression des caractères non imprimables et des caractères de contrôle, en conservant les caractères russes
    return re.sub(r'[^\x20-\x7Eа-яА-ЯёЁ]', '', text)

def save_text_to_docx(texts, docx_path_base, max_chars=980000):
    doc_num = 1
    char_count = 0
    
    doc = Document()
    print('save_text_to_docx: BEGIN')
    
    def create_new_doc():
        nonlocal doc, doc_num, char_count
        if char_count > 0:  # Save the current document if it has any content
            doc.save(f"{docx_path_base}_{doc_num:02d}.docx")
            doc_num += 1
            char_count = 0
            doc = Document()
        else:
            print('create_new_doc: NO CONTENT')
    
    for text, document_title in texts:
        # Ajouter le titre du document en gras
        title = doc.add_paragraph()
        title_run = title.add_run(document_title)
        title_run.bold = True
        title_run.font.size = Pt(14)
        char_count += len(document_title) + 2
        
        pages = text.split('Page ')
        for page in pages[1:]:  # Ignorer le premier élément car il est avant la première page
            page_content = page.split('\n\n', 1)
            page_num = page_content[0].strip()
            
            # Ajouter le numéro de page souligné
            page_paragraph = doc.add_paragraph()
            page_run = page_paragraph.add_run(f"Page {page_num}")
            page_run.underline = True
            char_count += len(page_run.text) + 2
            
            if len(page_content) > 1:
                paragraphs = page_content[1].split('\n\n')
                
                for para in paragraphs:
                    clean_para = clean_text(para)
                    if clean_para:  # Ajouter seulement si le paragraphe n'est pas vide
                        paragraph = doc.add_paragraph(clean_para)
                        char_count += len(clean_para) + 2
                    doc.add_paragraph('')  # Ajouter une ligne vide entre les paragraphes
                    char_count += 2
                    
                    if char_count >= max_chars:
                        create_new_doc()
    
    # Save the final document if it has any content
    if char_count > 0:
        doc.save(f"{docx_path_base}_{doc_num}.docx")
    else : 
        print('save_text_to_docx: NO CONTENT')

def process_directory(input_dir, output_dir):
    print('process_directory: BEGIN')
    texts = []
    output_base = os.path.join(output_dir, "document")

    for filename in os.listdir(input_dir):
        texts.clear()
        if filename.endswith(".pdf"):
            pdf_path = os.path.join(input_dir, filename)
            document_title = os.path.splitext(filename)[0]
            texte_extrait = extract_text_from_pdf(pdf_path)
            texts.append((texte_extrait, document_title))
    
            save_text_to_docx(texts, output_base + document_title)

def merge_docx_files(input_dir, output_path, max_chars=1000000):
    def get_docx_files(directory):
        return [os.path.join(directory, f) for f in os.listdir(directory) if f.endswith('.docx')]

    def create_composer(docx_files):
        base_doc = Document(docx_files[0])
        composer = Composer(base_doc)
        for docx_file in docx_files[1:]:
            composer.append(Document(docx_file))
        return composer

    docx_files = get_docx_files(input_dir)
    doc_num = 1
    char_count = 0
    current_composer = None

    for docx_file in docx_files:
        temp_doc = Document(docx_file)
        temp_text = '\n'.join([p.text for p in temp_doc.paragraphs])
        temp_char_count = len(temp_text)
        
        if current_composer is None:
            current_composer = create_composer([docx_file])
            char_count = temp_char_count
        else:
            if char_count + temp_char_count > max_chars:
                current_composer.save(f"{output_path}_{doc_num:02d}.docx")
                doc_num += 1
                current_composer = create_composer([docx_file])
                char_count = temp_char_count
            else:
                current_composer.append(Document(docx_file))
                char_count += temp_char_count
    
    if current_composer:
        current_composer.save(f"{output_path}_{doc_num:02d}.docx")

# Specify the input and output directory paths
input_dir = "/home/mtebad/projects/RAG-web-app/data/russian_data"
output_dir = "/home/mtebad/projects/RAG-web-app/data/output"

# changes pdfs to docx
process_directory(input_dir, output_dir)
# merges docx files
merge_docx_files(output_dir, os.path.join('/home/mtebad/projects/RAG-web-app/data/output', "merged_document"))

print("Traitement terminé avec succès.")
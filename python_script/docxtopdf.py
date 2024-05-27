from docx import Document
import os

def extract_and_save_segments(docx_path, output_dir):
    doc = Document(docx_path)
    current_segment = []
    current_title = None

    for paragraph in doc.paragraphs:
        if paragraph.runs and paragraph.runs[0].bold:
            # Sauvegarder le segment actuel
            if current_segment:
                if current_title:
                    save_segment_to_docx(current_segment, current_title, output_dir)
                current_segment = []
            current_title = paragraph.text.strip()  # Utiliser le texte en gras comme titre
        current_segment.append(paragraph)
    
    # Sauvegarder le dernier segment
    if current_segment and current_title:
        save_segment_to_docx(current_segment, current_title, output_dir)

def save_segment_to_docx(segment, title, output_dir):
    new_doc = Document()
    for paragraph in segment:
        new_paragraph = new_doc.add_paragraph()
        for run in paragraph.runs:
            new_run = new_paragraph.add_run(run.text)
            if run.bold:
                new_run.bold = True
            if run.italic:
                new_run.italic = True
            if run.underline:
                new_run.underline = True
            new_run.font.size = run.font.size
            new_run.font.name = run.font.name
    
    # Remplacer les caractères non valides pour les noms de fichiers
    valid_title = "".join(c if c.isalnum() else "_" for c in title)
    new_doc_path = os.path.join(output_dir, f"{valid_title}.docx")
    new_doc.save(new_doc_path)
    print(f"Saved: {new_doc_path}")

def process_directory(input_dir, output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    for filename in os.listdir(input_dir):
        if filename.endswith(".docx"):
            docx_path = os.path.join(input_dir, filename)
            extract_and_save_segments(docx_path, output_dir)

# Spécifiez les chemins du répertoire d'entrée et de sortie
input_dir = "russian_data_english_docx"
output_dir = "russian_data_english_pdf"
# Traiter tous les fichiers .docx dans le répertoire d'entrée
process_directory(input_dir, output_dir)

print("Traitement terminé avec succès.")



